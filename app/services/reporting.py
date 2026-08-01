"""Versioned, reproducible operational report compilation."""

from __future__ import annotations

import json as _json
from datetime import datetime, timezone
import io

from flask import current_app
from google.cloud import tasks_v2
from google.protobuf import timestamp_pb2
from docx import Document as WordDocument
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from app.database import db
from app.models.erp import BudgetLine, DocumentRecord, FeedbackForm, FeedbackResponse, ProjectSession, ReportSnapshot, SessionAttendance, WorkTask
from app.models.production import AggregateAttendance, ContributionRecord, ProjectRisk, ReportDefinition, ReportJob
from app.services.lifecycle import closure_blockers


def resolve_projects_in_scope(*, campus=None, program_type=None, academic_year=None, start_date=None, end_date=None):
    """Resolve the set of Project rows matching an optional
    campus/program-type/academic-year/date-range filter -- the same
    filter shape legacy campus-wide/program-wide reports used
    (`dashboard.py::_get_report_data`), lifted here as the shared entry
    point for multi-project rollup reports.
    """
    from app.models.project import Project

    query = Project.query
    if campus is not None:
        query = query.filter(Project.campus_id == campus.id)
    if program_type is not None:
        query = query.filter(Project.program_type_id == program_type.id)
    if academic_year is not None:
        query = query.filter(Project.academic_year_id == academic_year.id)
    if start_date is not None:
        query = query.filter(Project.end_date >= start_date)
    if end_date is not None:
        query = query.filter(Project.start_date <= end_date)
    return query.order_by(Project.id).all()


def _feedback_summary(project_ids):
    responses = (
        FeedbackResponse.query.join(FeedbackForm, FeedbackResponse.form_id == FeedbackForm.id)
        .filter(FeedbackForm.project_id.in_(project_ids), FeedbackResponse.moderation_status == "Approved")
        .all()
    )
    ratings = []
    for response in responses:
        answers = response.answers_json or {}
        rating = answers.get("rating") if isinstance(answers, dict) else None
        if isinstance(rating, (int, float)):
            ratings.append(float(rating))
    return {
        "response_count": len(responses),
        "average_rating": round(sum(ratings) / len(ratings), 2) if ratings else None,
    }


def _rollup_key(filters):
    relevant = {key: value for key, value in filters.items() if key != "_rollup_key"}
    return _json.dumps(relevant, sort_keys=True, default=str)


def compile_project_snapshot(project=None, *, projects=None, actor=None, report_type="Project Operational Report", filters=None):
    """Compile a versioned, approvable report snapshot for either a single
    project or a rollup across many (exactly one of `project`/`projects`
    must be given). A rollup snapshot has `project_id=None` and
    `source_references` listing every constituent project's public_id --
    this is the mechanism that replaces legacy campus-wide/program-wide
    reporting; there is no separate rollup model.
    """
    if bool(project) == bool(projects):
        raise ValueError("compile_project_snapshot requires exactly one of `project` or `projects`.")
    project_list = [project] if project else list(projects)
    project_ids = [p.id for p in project_list]

    individual_attendance = (
        SessionAttendance.query.join(ProjectSession)
        .filter(ProjectSession.project_id.in_(project_ids))
        .count()
    )
    aggregate_attendance = (
        db.session.query(db.func.coalesce(db.func.sum(AggregateAttendance.count), 0))
        .join(ProjectSession)
        .filter(ProjectSession.project_id.in_(project_ids))
        .scalar()
    )
    budget_totals = db.session.query(
        db.func.coalesce(db.func.sum(BudgetLine.estimated_amount), 0),
        db.func.coalesce(db.func.sum(BudgetLine.approved_amount), 0),
        db.func.coalesce(db.func.sum(BudgetLine.actual_amount), 0),
    ).filter(BudgetLine.project_id.in_(project_ids)).one()

    snapshot_data = {
        "reach": {
            "expected": sum(p.expected_reach or 0 for p in project_list),
            "actual": sum(p.actual_reach or 0 for p in project_list),
            "individual_attendance_records": individual_attendance,
            "aggregate_attendance": int(aggregate_attendance or 0),
        },
        "execution": {
            "tasks": WorkTask.query.filter(WorkTask.project_id.in_(project_ids)).count(),
            "documents": DocumentRecord.query.filter(DocumentRecord.project_id.in_(project_ids)).count(),
            "contributions": ContributionRecord.query.filter(ContributionRecord.project_id.in_(project_ids)).count(),
            "open_critical_risks": ProjectRisk.query.filter(
                ProjectRisk.project_id.in_(project_ids), ProjectRisk.status == "Open", ProjectRisk.is_critical == True
            ).count(),
        },
        "budget": {"estimated": str(budget_totals[0]), "approved": str(budget_totals[1]), "actual": str(budget_totals[2]), "currency": "INR"},
        "feedback": _feedback_summary(project_ids),
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }

    filters = dict(filters or {})
    if project:
        snapshot_data["project"] = {"public_id": project.public_id, "code": project.code, "title": project.title, "status": project.status}
        snapshot_data["execution"]["closure_blockers"] = [blocker.__dict__ for blocker in closure_blockers(project)]
        target_project_id = project.id
        source_references = [project.public_id]
        title = project.title
        next_version = (
            db.session.query(db.func.max(ReportSnapshot.version))
            .filter_by(project_id=target_project_id, report_type=report_type)
            .scalar() or 0
        ) + 1
    else:
        snapshot_data["projects"] = [
            {"public_id": p.public_id, "code": p.code, "title": p.title, "status": p.status} for p in project_list
        ]
        target_project_id = None
        source_references = [p.public_id for p in project_list]
        title = "Rollup report"
        rollup_key = _rollup_key(filters)
        filters["_rollup_key"] = rollup_key
        next_version = (
            db.session.query(db.func.max(ReportSnapshot.version))
            .filter(ReportSnapshot.project_id.is_(None), ReportSnapshot.report_type == report_type)
            .filter(ReportSnapshot.filters_json["_rollup_key"].as_string() == rollup_key)
            .scalar() or 0
        ) + 1

    snapshot = ReportSnapshot(
        project_id=target_project_id,
        report_type=report_type,
        title=f"{title} — {report_type}",
        version=next_version,
        filters_json=filters,
        snapshot_json=snapshot_data,
        source_references=source_references,
        approval_status="Draft",
        generated_by_id=getattr(actor, "id", None),
    )
    db.session.add(snapshot)
    db.session.flush()
    return snapshot


def execute_report_job(job):
    if job.status == "Completed":
        return job
    job.status = "Running"
    job.started_at = datetime.now(timezone.utc)
    try:
        project = db.session.get(__import__("app.models.project", fromlist=["Project"]).Project, job.project_id)
        if not project:
            raise ValueError("Report job has no accessible project.")
        definition = db.session.get(ReportDefinition, job.definition_id)
        snapshot = compile_project_snapshot(project, actor=db.session.get(__import__("app.models.user", fromlist=["User"]).User, job.requested_by_id), report_type=definition.report_type, filters=job.filters_json)
        job.snapshot_id = snapshot.id
        job.status = "Completed"
        job.completed_at = datetime.now(timezone.utc)
    except Exception as error:
        job.status = "Failed"
        job.error_summary = str(error)[:500]
        job.completed_at = datetime.now(timezone.utc)
    db.session.commit()
    return job


def enqueue_report_job(job):
    if current_app.config.get("APP_ENV") != "production":
        return None
    client = tasks_v2.CloudTasksClient()
    parent = client.queue_path(current_app.config["GCP_PROJECT_ID"], current_app.config["GCP_REGION"], current_app.config["CLOUD_TASKS_QUEUE"])
    schedule_time = timestamp_pb2.Timestamp()
    schedule_time.FromDatetime(datetime.now(timezone.utc))
    task = {
        "name": f"{parent}/tasks/report-{job.public_id}",
        "schedule_time": schedule_time,
        "http_request": {
            "http_method": tasks_v2.HttpMethod.POST,
            "url": f"{current_app.config['INTERNAL_JOB_BASE_URL'].rstrip('/')}/internal/jobs/reports/{job.public_id}/execute",
            "oidc_token": {
                "service_account_email": current_app.config["TASKS_SERVICE_ACCOUNT"],
                "audience": current_app.config["INTERNAL_JOB_AUDIENCE"],
            },
        },
    }
    try:
        return client.create_task(parent=parent, task=task)
    except Exception:
        current_app.logger.exception("Unable to enqueue report job %s", job.public_id)
        return None


def _flatten(prefix, value, rows):
    if isinstance(value, dict):
        for key, nested in value.items():
            _flatten(f"{prefix}.{key}" if prefix else key, nested, rows)
    elif isinstance(value, list) and value and all(isinstance(item, dict) for item in value):
        # A list of dicts (e.g. a rollup snapshot's "projects" breakdown)
        # renders as its own indexed sub-table rather than being squashed
        # into one unreadable joined string.
        for index, item in enumerate(value, start=1):
            _flatten(f"{prefix}[{index}]" if prefix else f"[{index}]", item, rows)
    elif isinstance(value, list):
        rows.append((prefix, "; ".join(str(item) for item in value)))
    else:
        rows.append((prefix, "" if value is None else str(value)))


def render_snapshot(snapshot, output_format):
    rows = []
    _flatten("", snapshot.snapshot_json, rows)
    output = io.BytesIO()
    if output_format == "xlsx":
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Report"
        sheet.append([snapshot.title, f"Version {snapshot.version}", snapshot.approval_status])
        sheet.append(["Field", "Value"])
        for row in rows:
            sheet.append(row)
        sheet.freeze_panes = "A3"
        workbook.save(output)
        mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    elif output_format == "docx":
        document = WordDocument()
        document.add_heading(snapshot.title, 0)
        document.add_paragraph(f"Version {snapshot.version} · {snapshot.approval_status} · {snapshot.created_at}")
        table = document.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "Field"
        table.rows[0].cells[1].text = "Value"
        for field, value in rows:
            cells = table.add_row().cells
            cells[0].text = field
            cells[1].text = value
        document.save(output)
        mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    elif output_format == "pdf":
        pdf = canvas.Canvas(output, pagesize=A4)
        width, height = A4
        y = height - 48
        pdf.setTitle(snapshot.title)
        pdf.setFont("Helvetica-Bold", 14)
        pdf.drawString(40, y, snapshot.title[:90])
        y -= 24
        pdf.setFont("Helvetica", 8)
        for field, value in rows:
            line = f"{field}: {value}"[:150]
            if y < 48:
                pdf.showPage()
                pdf.setFont("Helvetica", 8)
                y = height - 48
            pdf.drawString(40, y, line)
            y -= 12
        pdf.save()
        mime_type = "application/pdf"
    else:
        raise ValueError("Supported report formats are xlsx, docx, and pdf.")
    output.seek(0)
    return output, mime_type
